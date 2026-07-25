"""DOCX-Dateien einlesen."""

from docx import Document as WordDocument

from app.core.document import Document
from app.logger import logger

class DocumentReader:
    """Liest ein Word-Dokument ein."""

    def read(self, filename: str) -> Document:
        word = WordDocument(filename)

        document = Document()

        for paragraph in word.paragraphs:
            document.add_paragraph(paragraph.text)

        logger.info(
            "Dokument erfolgreich eingelesen (%d Absätze)",
            document.paragraph_count,
        )


        return document