"""DOCX-Dateien schreiben."""

from docx import Document as WordDocument

from app.core.document import Document
from app.logger import logger

class DocumentWriter:
    """Schreibt ein Word-Dokument."""

    def write(self, document: Document, filename: str) -> None:
        word = WordDocument()

        for paragraph in document.paragraphs:
            word.add_paragraph(paragraph.text)

        word.save(filename)
        logger.info(
            "Dokument erfolgreich gespeichert (%d Absätze)",
            document.paragraph_count,
        )